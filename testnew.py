import markdown
import weasyprint
from docx import Document
from bs4 import BeautifulSoup


def markdown_to_html(md_content):
    """Convert Markdown to HTML."""
    return markdown.markdown(md_content, extensions=['tables'])


def save_as_pdf(html_content, output_pdf):
    """Convert HTML content to PDF using WeasyPrint."""
    html_template = f"""
    <html>
    <head>
        <style>
            table {{
                width: 100%;
                border-collapse: collapse;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #f2f2f2;
            }}
            body {{
                font-family: Arial, sans-serif;
                margin: 20px;
            }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    weasyprint.HTML(string=html_template).write_pdf(output_pdf)
    print(f"PDF saved as {output_pdf}")


def save_as_word(html_content, output_docx):
    """Convert HTML content to Word document using python-docx."""
    doc = Document()

    # Parse the HTML content
    soup = BeautifulSoup(html_content, "html.parser")

    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'li', 'table']):
        if element.name in ['h1', 'h2', 'h3']:
            doc.add_heading(element.text, level=int(element.name[1]))
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style="ListBullet")  # FIX: Removed manual "• "
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.text

    doc.save(output_docx)
    print(f"Word document saved as {output_docx}")


def convert_markdown(markdown_file, output_pdf, output_docx):
    """Convert Markdown file to both PDF and Word."""
    with open(markdown_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    html_content = markdown_to_html(md_content)
    save_as_pdf(html_content, output_pdf)
    save_as_word(html_content, output_docx)


# File paths
markdown_file = 'test2.md'
output_pdf = 'output.pdf'
output_docx = 'output.docx'

convert_markdown(markdown_file, output_pdf, output_docx)
